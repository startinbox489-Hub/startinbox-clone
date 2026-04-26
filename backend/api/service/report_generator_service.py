"""
Reposrt generator service
"""

import os
from typing import Any, Dict

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as PDFImage,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from uuid6 import uuid7

from api.utils.task_logger import create_logger

logger = create_logger(":: REPORT GENERATOR ::")

body_wrap_style = ParagraphStyle(
    name="BodyWrap",
    fontSize=10,
    leading=14,
    spaceAfter=4,
    wordWrap="CJK",  # handles long words / URLs
)

metric_style = ParagraphStyle(
    name="Metric",
    fontSize=16,
    leading=20,
    textColor=colors.HexColor("#470F86"),
    spaceAfter=12,
)


def set_cell_margins(cell, top=100, bottom=100, start=100, end=100):
    """
    Margins are in twips (1/20 pt). 100 ≈ 5pt
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")

    for tag, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)

    tcPr.append(tcMar)


def keep_row_together(row):
    """
    Keep table rows together

    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def section_heading(doc, text, level=1):
    """
    section heading

    :param doc: Description
    :param text: Description
    :param level: Description
    """
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h


def keep_cell_paragraphs_together(cell):
    """
    keep_cell_paragraphs_together

    :param cell: Description
    """
    for paragraph in cell.paragraphs:
        pPr = paragraph._p.get_or_add_pPr()

        keep_next = OxmlElement("w:keepNext")
        keep_lines = OxmlElement("w:keepLines")

        pPr.append(keep_next)
        pPr.append(keep_lines)


class PDFError(Exception):
    """
    PDF error
    """

    def __init__(self, message: str = "Error generating PDF") -> None:
        """
        Constructor
        """
        super().__init__(message)


class DocxError(Exception):
    """
    DOCX error
    """

    def __init__(self, message: str = "Error generating DOCX") -> None:
        """
        Constructor
        """
        super().__init__(message)


class ReportGeneratorService:
    """
    ReportGenerator Service
    """

    def __init__(self) -> None:
        """
        Constructor
        """

        self.__cwd = os.getcwd()
        self.__logo_path = f"{self.__cwd}/assets/logo-png-1.png"

    # ================= PDF EXPORT ================= #
    def export_pdf(
        self,
        data: dict[str, Any],
        filename: str = f"{str(uuid7())}_startup_report.pdf",
        validated_idea: str | None = None,
    ) -> str:
        """
        Export pdf.

        Args:
            filename (str): The name of the file.
        Returns:
            str: The filename
        """
        try:

            file_path = f"{self.__cwd}/startup_reports/pdfs/{filename}"

            doc = SimpleDocTemplate(file_path)

            styles = getSampleStyleSheet()

            elements = []

            # Logo
            try:

                elements.append(PDFImage(self.__logo_path, width=120, height=60))

                elements.append(Spacer(1, 20))
            except (
                ValueError,
                IOError,
                FileExistsError,
                FileNotFoundError,
                OSError,
            ) as exc:

                logger.error(msg=f"error getting image for pdf export {str(exc)}")

            # ==================== SECTION A ==================== #

            elements.append(
                Paragraph("🚀 Startup Idea Validation Report", styles["Title"])
            )

            elements.append(Spacer(1, 20))
            a_validation_output: dict[str, Any] | None = data.get(
                "A. Validation Output"
            )
            if validated_idea:

                elements.append(
                    Paragraph(
                        f"<b>Validated Idea:</b> <font size='15'>{validated_idea}</font>",
                        metric_style,
                    )
                )

                elements.append(Spacer(1, 20))
            if a_validation_output:
                # Idea Score

                score = a_validation_output.get("idea_score")

                elements.append(
                    Paragraph(
                        f"<b>Idea Score:</b> <font size='10'>{score}/100</font>",
                        metric_style,
                    )
                )

                elements.append(Spacer(1, 20))

                # Lean Canvas
                elements.append(Paragraph("Lean Canvas", styles["Heading1"]))

                lean: dict[str, Any] = a_validation_output.get("lean_canvas") or {}

                table_data = [["Block", "Description"]] + [
                    [
                        Paragraph(k.replace("_", " ").title(), body_wrap_style),
                        Paragraph(str(v), body_wrap_style),
                    ]
                    for k, v in lean.items()
                ]

                table = Table(table_data, colWidths=[150, 350])

                table.setStyle(
                    TableStyle(
                        [
                            # Header styling
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E86AB")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            # Header styling
                            ("BOX", (0, 0), (-1, -1), 1, colors.black),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            # readable wrapping
                            ("LEFTPADDING", (0, 0), (-1, -1), 8),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                            # Good layout
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )

                elements.append(table)

                elements.append(Spacer(1, 20))

                # Persona
                elements.append(Paragraph("Ideal Customer Persona", styles["Heading1"]))

                persona: dict[str, Any] = (
                    a_validation_output.get("ideal_customer_persona") or {}
                )

                persona_table = [["Attribute", "Details"]] + [
                    [
                        Paragraph(k.replace("_", " ").title(), body_wrap_style),
                        Paragraph(str(v), body_wrap_style),
                    ]
                    for k, v in persona.items()
                ]

                table = Table(persona_table, colWidths=[150, 350])

                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E86AB")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("BOX", (0, 0), (-1, -1), 1, colors.black),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            # readable wrapping
                            ("LEFTPADDING", (0, 0), (-1, -1), 8),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                            # Good layout
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )

                elements.append(table)

                elements.append(Spacer(1, 20))

                # suggested startup names
                suggested_startup_names: dict[str, Any] = (
                    a_validation_output.get("suggested_startup_names") or {}
                )
                if suggested_startup_names:
                    elements.append(
                        Paragraph("Suggested Startup Names", styles["Heading1"])
                    )

                    persona_table = [["Attribute", "Details"]] + [
                        [
                            Paragraph(str(idx + 1), body_wrap_style),
                            Paragraph(v, body_wrap_style),
                        ]
                        for idx, v in enumerate(suggested_startup_names)
                    ]

                    table = Table(persona_table, colWidths=[150, 350])

                    table.setStyle(
                        TableStyle(
                            [
                                (
                                    "BACKGROUND",
                                    (0, 0),
                                    (-1, 0),
                                    colors.HexColor("#2E86AB"),
                                ),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                                # readable wrapping
                                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                                ("TOPPADDING", (0, 0), (-1, -1), 6),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                                # Good layout
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ]
                        )
                    )

                    elements.append(table)

                    elements.append(Spacer(1, 20))

                # monetization models
                monetization_models: dict[str, Any] = (
                    a_validation_output.get("monetization_models") or {}
                )
                if monetization_models:
                    elements.append(
                        Paragraph("Monetization Models", styles["Heading1"])
                    )

                    persona_table = [["Attribute", "Details"]] + [
                        [
                            Paragraph(str(idx + 1), body_wrap_style),
                            Paragraph(v, body_wrap_style),
                        ]
                        for idx, v in enumerate(monetization_models)
                    ]

                    table = Table(persona_table, colWidths=[150, 350])

                    table.setStyle(
                        TableStyle(
                            [
                                (
                                    "BACKGROUND",
                                    (0, 0),
                                    (-1, 0),
                                    colors.HexColor("#2E86AB"),
                                ),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                                # readable wrapping
                                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                                ("TOPPADDING", (0, 0), (-1, -1), 6),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                                # Good layout
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ]
                        )
                    )

                    elements.append(table)

                    # elements.append(PageBreak())
                    elements.append(Spacer(1, 20))

            # ==================== SECTION B ==================== #
            # Website Hero
            b_launch_content_generator: dict[str, Any] | None = (
                data.get("B. Launch Content Generator") or {}
            )
            if b_launch_content_generator:
                hero: dict[str, Any] = (
                    b_launch_content_generator.get("website_hero") or {}
                )

                elements.append(Paragraph("Website Hero Section", styles["Heading1"]))

                elements.append(
                    Paragraph(f"<b>{hero.get('headline','')}</b>", styles["Heading2"])
                )

                elements.append(
                    Paragraph(hero.get("subheadline", ""), styles["Normal"])
                )

                for f in hero.get("features", []):

                    elements.append(Paragraph(f"• {f}", styles["Normal"]))

                # elements.append(PageBreak())
                elements.append(Spacer(1, 20))

                # Blog Posts
                elements.append(Paragraph("Blog Post Ideas", styles["Heading1"]))

                for post in b_launch_content_generator.get("blog_posts", {}):

                    elements.append(
                        Paragraph(f"<b>{post.get('title')}</b>", styles["Normal"])
                    )

                    for bullet in post.get("outline"):

                        elements.append(Paragraph(f"• {bullet}", styles["Normal"]))

                    elements.append(Spacer(1, 10))

                # Social Posts
                elements.append(
                    Paragraph("Social Media Launch Posts", styles["Heading1"])
                )

                for tw in b_launch_content_generator.get("twitter_posts", []):

                    elements.append(Paragraph(f"• {tw}", styles["Normal"]))

            # ==================== SECTION C ==================== #

            infleuncer_outreach: dict[str, Any] = data.get(
                "C. Influencer Outreach Generator", {}
            )

            if infleuncer_outreach:

                influencers = infleuncer_outreach.get("influencers", [])

                if influencers:

                    # elements.append(PageBreak())
                    elements.append(Spacer(1, 20))

                    elements.append(
                        Paragraph("Influencer Outreach Generator", styles["Heading1"])
                    )

                    table_data = [
                        [
                            Paragraph("Handle", body_wrap_style),
                            Paragraph("Niche", body_wrap_style),
                            Paragraph("Outreach DM", body_wrap_style),
                        ]
                    ]

                    for inf in influencers:

                        table_data.append(
                            [
                                Paragraph(inf["handle"], body_wrap_style),
                                Paragraph(inf["niche"], body_wrap_style),
                                Paragraph(inf["dm"], body_wrap_style),
                            ]
                        )

                    table = Table(table_data, colWidths=[120, 120, 260])

                    table.setStyle(
                        TableStyle(
                            [
                                (
                                    "BACKGROUND",
                                    (0, 0),
                                    (-1, 0),
                                    colors.HexColor("#2E86AB"),
                                ),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                # readable wrapping
                                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                                ("TOPPADDING", (0, 0), (-1, -1), 6),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                                # Good layout
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                #  prevents row break
                                ("NOSPLIT", (0, 0), (-1, -1)),
                            ]
                        )
                    )

                    elements.append(table)

            doc.build(elements)

            logger.info("PDF saved in %s", file_path)

            return filename
        except Exception as exc:

            logger.error("Error generating pdf: %s", str(exc))
            raise PDFError() from exc

    # ================= DOCX EXPORT ================= #
    def export_docx(
        self,
        data: dict[str, Any],
        filename: str = f"{str(uuid7())}_startup_report.docx",
        validated_idea: str | None = None,
    ) -> str:
        """
        Eport docx.

        Args:
            filename (str): The name of the file.
        Returns:
            str: The filename
        """
        try:
            doc = Document()
            file_path = f"{self.__cwd}/startup_reports/docs/{filename}"

            # Logo
            try:
                doc.add_picture(self.__logo_path, width=Inches(1.5))
            except (
                ValueError,
                IOError,
                FileExistsError,
                FileNotFoundError,
            ) as exc:
                logger.error(msg=f"error getting image for pdf export {str(exc)}")

            # ==================== SECTION A ==================== #
            doc.add_heading("🚀 Startup Idea Validation Report", 0)
            doc.add_paragraph()
            doc.add_paragraph()
            if validated_idea:
                doc.add_paragraph(f"Validated Idea: {validated_idea}")
                doc.add_paragraph()
            score = data.get("A. Validation Output", {}).get("idea_score")

            p = doc.add_paragraph()
            run = p.add_run(f"Idea Score: {score}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

            doc.add_page_break()
            a_validation_output = data.get("A. Validation Output") or {}

            if a_validation_output:

                # Lean Canvas
                doc.add_heading("Lean Canvas", level=1)
                lean: dict[str, Any] = a_validation_output.get("lean_canvas") or {}
                table = doc.add_table(rows=1, cols=2)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text = "Block", "Description"
                for k, v in lean.items():
                    row = table.add_row().cells
                    for cell in row:
                        set_cell_margins(cell)
                    row[0].text = k.capitalize()
                    row[1].text = v
                doc.add_paragraph()

                # Persona
                doc.add_heading("Ideal Customer Persona", level=1)
                persona: dict[str, Any] = (
                    a_validation_output.get("ideal_customer_persona") or {}
                )
                table = doc.add_table(rows=1, cols=2)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text = "Attribute", "Details"
                for k, v in persona.items():
                    row = table.add_row().cells
                    for cell in row:
                        set_cell_margins(cell)
                    row[0].text = k.capitalize()
                    row[1].text = v
                doc.add_page_break()

                # Startup names
                suggested_startup_names: dict[str, Any] = (
                    a_validation_output.get("suggested_startup_names") or {}
                )
                if suggested_startup_names:
                    doc.add_heading("Suggested Startup Names", level=1)
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text = "Attribute", "Details"
                    for idx, v in enumerate(suggested_startup_names):
                        row = table.add_row().cells
                        for cell in row:
                            set_cell_margins(cell)
                        row[0].text = str(idx + 1)
                        row[1].text = v
                    doc.add_page_break()

                # Monetization models
                monetization_models: dict[str, Any] = (
                    a_validation_output.get("monetization_models") or {}
                )
                if monetization_models:
                    doc.add_heading("Monetization Models", level=1)
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text = "Attribute", "Details"
                    for idx, v in enumerate(monetization_models):
                        row = table.add_row().cells
                        for cell in row:
                            set_cell_margins(cell)
                        row[0].text = str(idx + 1)
                        row[1].text = v
                    doc.add_page_break()

            # ==================== SECTION B ==================== #
            # Website Hero
            b_launch_content_generator = data.get("B. Launch Content Generator") or {}
            if b_launch_content_generator:
                hero = b_launch_content_generator["website_hero"]
                # doc.add_heading("Website Hero Section", level=1)
                section_heading(doc, "Website Hero Section")
                doc.add_paragraph(hero["headline"]).runs[0].bold = True
                doc.add_paragraph(hero["subheadline"])
                doc.add_heading("Features:", level=2)
                for f in hero["features"]:
                    doc.add_paragraph(f, style="List Bullet")

                # Blog Posts
                doc.add_heading("Blog Post Ideas", level=1)
                for post in b_launch_content_generator["blog_posts"]:
                    doc.add_paragraph(post["title"]).runs[0].bold = True
                    for bullet in post["outline"]:
                        doc.add_paragraph(bullet, style="List Bullet")

                # Social Posts
                doc.add_heading("Social Media Launch Posts", level=1)
                for tw in b_launch_content_generator.get("twitter_posts", []):
                    doc.add_paragraph(tw, style="List Bullet")

                # Elevator Pitch
                pitch = b_launch_content_generator["elevator_pitch_slide"]
                doc.add_heading("Elevator Pitch Slide", level=1)
                doc.add_paragraph(pitch["headline"]).runs[0].bold = True
                for bp in pitch["bullet_points"]:
                    doc.add_paragraph(bp, style="List Bullet")

            # ==================== SECTION C ==================== #
            infleuncer_outreach: dict[str, Any] = data.get(
                "C. Influencer Outreach Generator", {}
            )
            if infleuncer_outreach:

                influencers = infleuncer_outreach.get("influencers", [])
                if influencers:
                    doc.add_page_break()
                    doc.add_heading("Influencer Outreach Generator", level=1)

                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Light Grid Accent 1"
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text = (
                        "Handle",
                        "Niche",
                        "Outreach DM",
                    )

                    for inf in influencers:
                        row = table.add_row().cells
                        for cell in row:
                            set_cell_margins(cell)
                            keep_cell_paragraphs_together(cell)
                        row[0].text = inf["handle"]
                        row[1].text = inf["niche"]
                        row[2].text = inf["dm"]

            doc.save(file_path)
            logger.info("DOCX saved in %s", file_path)
            return filename

        except Exception as exc:
            logger.error("Error generating docx: %s", str(exc))
            raise DocxError() from exc

    def prep_for_file_gen(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Transforms raw AI output into a structured, file-generation-ready format.

        Sections:
        A. Validation Output
        B. Launch Content Generator
        C. Influencer Outreach Generator
        """
        try:
            if data.pop("_sa_instance_state", None):
                data.pop("_sa_instance_state", None)
            if data.pop("created_at", None):
                data.pop("created_at", None)
            if data.pop("updated_at", None):
                data.pop("updated_at", None)
            if data.pop("id", None):
                data.pop("id", None)
            if data.pop("user_id", None):
                data.pop("user_id", None)

            return_value: Dict[str, Any] = {
                "A. Validation Output": {},
                "B. Launch Content Generator": {},
                "C. Influencer Outreach Generator": {},
            }
            # A. Validation Output
            if data.get("idea_validation"):
                return_value["A. Validation Output"]["idea_validation"] = data[
                    "idea_validation"
                ]

            if data.get("idea_score") is not None:
                return_value["A. Validation Output"]["idea_score"] = data["idea_score"]

            if data.get("lean_canvas"):
                return_value["A. Validation Output"]["lean_canvas"] = data[
                    "lean_canvas"
                ]

            if data.get("ideal_customer_persona"):
                return_value["A. Validation Output"]["ideal_customer_persona"] = data[
                    "ideal_customer_persona"
                ]

            if data.get("suggested_startup_names"):
                return_value["A. Validation Output"]["suggested_startup_names"] = data[
                    "suggested_startup_names"
                ]

            if data.get("monetization_models"):
                return_value["A. Validation Output"]["monetization_models"] = data[
                    "monetization_models"
                ]
            # B. Launch Content Generator
            if data.get("website_hero"):
                return_value.update({})
                return_value["B. Launch Content Generator"]["website_hero"] = data[
                    "website_hero"
                ]

            if data.get("blog_posts"):
                return_value["B. Launch Content Generator"]["blog_posts"] = data[
                    "blog_posts"
                ]

            if data.get("twitter_posts"):
                return_value["B. Launch Content Generator"]["twitter_posts"] = data[
                    "twitter_posts"
                ]

            if data.get("elevator_pitch_slide"):
                return_value["B. Launch Content Generator"]["elevator_pitch_slide"] = (
                    data["elevator_pitch_slide"]
                )

            # C. Influencer Outreach Generator
            influencers = []

            if data.get("influencer_one"):
                influencers.append(data["influencer_one"])

            if data.get("influencer_two"):
                influencers.append(data["influencer_two"])

            if data.get("influencer_three"):
                influencers.append(data["influencer_three"])

            if influencers:
                return_value["C. Influencer Outreach Generator"][
                    "influencers"
                ] = influencers

            if data.get("go_to_market_strategy_outline"):
                return_value["C. Influencer Outreach Generator"][
                    "go_to_market_strategy_outline"
                ] = data["go_to_market_strategy_outline"]

            return {
                section: content for section, content in return_value.items() if content
            }
        except Exception as exc:
            print(exc)
            raise exc


report_generator_service = ReportGeneratorService()
