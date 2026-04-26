"""
DOCSX Service
"""

from typing import Any, Dict
from docx import Document
from docx.shared import RGBColor


class DOCSXService:
    """
    DOCSX Service
    """

    def __init__(self) -> None:
        """
        Constructor
        """

        return

    def export_to_docx(
        self, data: Dict[str, Any], filename: str = "startup_report.docx"
    ):
        """
        Formats and exports to  DOCX
        """
        doc = Document()
        doc.add_heading("🚀 Startup Idea Validation Report", 0)

        # Idea Score
        doc.add_heading("Idea Score", level=1)
        doc.add_paragraph(str(data["A. Validation Output"]["idea_score"]))

        # Lean Canvas
        doc.add_heading("Lean Canvas", level=1)
        for key, value in data["A. Validation Output"]["lean_canvas"].items():
            doc.add_paragraph(f"{key}: {value}")

        # Customer Persona
        doc.add_heading("Ideal Customer Persona", level=1)
        for key, value in data["A. Validation Output"][
            "ideal_customer_persona"
        ].items():
            doc.add_paragraph(f"{key}: {value}")

        # Suggested Names
        doc.add_heading("Suggested Startup Names", level=1)
        for name in data["A. Validation Output"]["suggested_startup_names"]:
            doc.add_paragraph(f"- {name}")

        # Launch Content
        doc.add_heading("Launch Content", level=1)
        hero = data["B. Launch Content Generator"]["website_hero"]
        doc.add_paragraph(f"Website Hero: {hero['headline']} - {hero['subheadline']}")
        for f in hero["features"]:
            doc.add_paragraph(f"- {f}")

        doc.save(filename)
        print(f"DOCX saved as {filename}")

    def export_polished_docx(
        self, data: Dict[str, Any], filename="startup_report.docx"
    ):
        """
        Polishes, Formats, and Exports docsx
        """
        doc = Document()

        # === TITLE PAGE ===
        title = doc.add_heading("🚀 Startup Idea Validation Report", 0)
        run = title.runs[0]
        run.font.color.rgb = RGBColor(46, 134, 171)  # #2E86AB
        doc.add_paragraph(f"Idea Score: {data['A. Validation Output']['idea_score']}")

        doc.add_page_break()

        # === LEAN CANVAS (TABLE) ===
        doc.add_heading("Lean Canvas", level=1)
        lean = data["A. Validation Output"]["lean_canvas"]
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Block"
        hdr_cells[1].text = "Description"
        for k, v in lean.items():
            row_cells = table.add_row().cells
            row_cells[0].text = k.capitalize()
            row_cells[1].text = v
        doc.add_paragraph()

        # === CUSTOMER PERSONA ===
        doc.add_heading("Ideal Customer Persona", level=1)
        persona = data["A. Validation Output"]["ideal_customer_persona"]
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Attribute"
        hdr_cells[1].text = "Details"
        for k, v in persona.items():
            row_cells = table.add_row().cells
            row_cells[0].text = k.capitalize()
            row_cells[1].text = v
        doc.add_paragraph()

        # === STARTUP NAMES ===
        doc.add_heading("Suggested Startup Names", level=1)
        for name in data["A. Validation Output"]["suggested_startup_names"]:
            doc.add_paragraph(name, style="List Bullet")

        # === MONETIZATION MODELS ===
        doc.add_heading("Monetization Models", level=1)
        for model in data["A. Validation Output"]["monetization_models"]:
            doc.add_paragraph(model, style="List Bullet")

        doc.add_page_break()

        # === WEBSITE HERO ===
        doc.add_heading("Website Hero Section", level=1)
        hero = data["B. Launch Content Generator"]["website_hero"]
        doc.add_paragraph(hero["headline"]).runs[0].bold = True
        doc.add_paragraph(hero["subheadline"])
        doc.add_heading("Features:", level=2)
        for f in hero["features"]:
            doc.add_paragraph(f, style="List Bullet")

        # === BLOG POSTS ===
        doc.add_heading("Blog Post Ideas", level=1)
        for post in data["B. Launch Content Generator"]["blog_posts"]:
            doc.add_paragraph(post["title"]).runs[0].bold = True
            for bullet in post["outline"]:
                doc.add_paragraph(bullet, style="List Bullet")

        doc.add_page_break()

        # === TWITTER POSTS ===
        doc.add_heading("Social Media Launch Posts", level=1)
        for tw in data["B. Launch Content Generator"]["twitter_posts"]:
            doc.add_paragraph(tw, style="List Bullet")

        # === ELEVATOR PITCH ===
        doc.add_heading("Elevator Pitch Slide", level=1)
        pitch = data["B. Launch Content Generator"]["elevator_pitch_slide"]
        doc.add_paragraph(pitch["headline"]).runs[0].bold = True
        for bp in pitch["bullet_points"]:
            doc.add_paragraph(bp, style="List Bullet")

        # Save file
        doc.save(filename)
        print(f"✅ Polished DOCX saved as {filename}")


docx_service = DOCSXService()
